# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2026 SourceBridge Contributors

"""Word (.docx) renderer.

Converts Markdown to a Word document with proper styles so users
can edit the report and update the TOC. Uses python-docx.

python-docx is optional — if not installed, Word rendering is skipped.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)


def render_docx(markdown_path: str, docx_path: str) -> bool:
    """Render a Markdown file to a Word document.

    Returns True if successful, False if python-docx is not available.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("python-docx not installed — Word rendering skipped. Install with: pip install python-docx")
        return False

    md_content = Path(markdown_path).read_text()
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Parse Markdown line by line
    lines = md_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            doc.add_paragraph("_" * 50)
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else "Normal"
            p.paragraph_format.left_indent = Inches(0.5)
            _add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)
        i += 1

    try:
        doc.save(docx_path)
        logger.info("docx_rendered", docx_path=docx_path)
        return True
    except Exception as e:
        logger.error("docx_render_failed", error=str(e))
        return False


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text to a paragraph with basic Markdown formatting (bold, italic, code)."""
    # Simple regex-based formatting
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            try:
                from docx.shared import Pt
                run.font.size = Pt(9)
            except ImportError:
                pass
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
